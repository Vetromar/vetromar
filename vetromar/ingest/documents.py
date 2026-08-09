"""Local document parsing: PDF/DOCX/MD/TXT file -> episode raw text + a
structural locator table.

The parsed text becomes `episode.raw` — the evidence gate then works exactly
as for any text source (excerpts must be literal spans). What documents add
is the locator table: each page/paragraph's offsets within the NORMALIZED
text (the same whitespace normalization the gate applies), so excerpt
evidence can be stamped with a human-usable position ("page:4") in the
existing free-text `ExcerptEvidence.locator` field — no schema change, and
the frozen surfaces are nowhere near this path.

Parsing is local-only (pypdf / python-docx, pure Python). Layout fidelity is
deliberately not a goal: evidence is literal text spans, not bounding boxes.
"""

from __future__ import annotations

from bisect import bisect_right
from dataclasses import dataclass
from pathlib import Path

from vetromar.errors import ConfigError
from vetromar.extraction.validate import _normalize

SUPPORTED_SUFFIXES = (".pdf", ".docx", ".md", ".txt")


@dataclass
class ParsedDocument:
    text: str
    # (start, end, locator) over the NORMALIZED text, e.g. (0, 812, "page:1").
    spans: list[tuple[int, int, str]]


def parse_document(path: Path) -> ParsedDocument:
    """Extract text + locator spans from a document file. Raises ConfigError
    for unsupported types, unreadable files, and missing parser deps."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        segments = _pdf_segments(path)
    elif suffix == ".docx":
        segments = _docx_segments(path)
    elif suffix in (".md", ".txt"):
        raw = path.read_text(errors="replace")
        paragraphs = [p for p in raw.split("\n\n") if p.strip()]
        segments = [(p, f"para:{i + 1}") for i, p in enumerate(paragraphs)]
    else:
        raise ConfigError(
            f"Unsupported document type: {path.suffix or path.name}",
            hint=f"Supported: {', '.join(SUPPORTED_SUFFIXES)}",
        )
    text_parts: list[str] = []
    spans: list[tuple[int, int, str]] = []
    cursor = 0
    for segment_text, locator in segments:
        norm = _normalize(segment_text)
        if not norm:
            continue
        start = cursor if not spans else cursor + 1  # the " " joiner
        spans.append((start, start + len(norm), locator))
        cursor = start + len(norm)
        text_parts.append(segment_text)
    if not spans:
        raise ConfigError(
            f"No extractable text in {path.name}.",
            hint="Scanned/image-only PDFs need OCR first — Vetromar reads text layers.",
        )
    return ParsedDocument(text="\n\n".join(text_parts), spans=spans)


def _pdf_segments(path: Path) -> list[tuple[str, str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ConfigError(
            "PDF support isn't installed.",
            hint='pip install "vetromar[documents]" (bundled in the desktop app)',
        ) from exc
    try:
        reader = PdfReader(str(path))
        return [
            (page.extract_text() or "", f"page:{i + 1}")
            for i, page in enumerate(reader.pages)
        ]
    except ConfigError:
        raise
    except Exception as exc:  # noqa: BLE001 — malformed uploads must fail friendly
        raise ConfigError(
            f"Could not read {path.name} as a PDF.",
            hint="Is the file password-protected or corrupted?",
        ) from exc


def _docx_segments(path: Path) -> list[tuple[str, str]]:
    try:
        import docx
    except ImportError as exc:
        raise ConfigError(
            "DOCX support isn't installed.",
            hint='pip install "vetromar[documents]" (bundled in the desktop app)',
        ) from exc
    try:
        document = docx.Document(str(path))
        return [
            (p.text, f"para:{i + 1}") for i, p in enumerate(document.paragraphs)
        ]
    except Exception as exc:  # noqa: BLE001
        raise ConfigError(
            f"Could not read {path.name} as a DOCX document.",
            hint="Is the file corrupted, or an old .doc? Only .docx is supported.",
        ) from exc


def locator_for(parsed: ParsedDocument, norm_position: int) -> str | None:
    """The locator of the span containing a position in the normalized text
    (an excerpt straddling a boundary reports where it starts)."""
    starts = [start for start, _end, _loc in parsed.spans]
    idx = bisect_right(starts, norm_position) - 1
    if idx < 0:
        return None
    start, end, locator = parsed.spans[idx]
    return locator if norm_position < end else None


def stamp_locators(drafts, parsed: ParsedDocument) -> None:
    """Fill each excerpt's empty `locator` with its structural position.
    Runs AFTER quote healing (texts are final and literal); evidence text is
    never modified — only the locator annotation is added."""
    norm_full = _normalize(parsed.text)
    for draft in drafts:
        for ev in draft.evidence:
            if ev.kind != "excerpt" or getattr(ev, "locator", None):
                continue
            position = norm_full.find(_normalize(ev.text))
            if position < 0:
                continue  # gate will judge it; locators are best-effort
            ev.locator = locator_for(parsed, position)
